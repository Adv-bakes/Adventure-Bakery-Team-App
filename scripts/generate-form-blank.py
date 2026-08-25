"""Schema-driven blank printable forms (PDF + DOCX) for the consolidated sanitation forms.
Reads the committed form_schema JSON, lays out each section/field in the Adventure Bakery
house style (reportlab + python-docx). Run from repo root.

Blank = handwrite on paper -> photograph -> enter into the digital form.
- FRM-903 (one record per production day): one sheet = one day, all sections rendered once.
- FRM-901 (periodic): schedule reference table (printed) + a multi-row completion LOG
  (each row -> one digital entry).
"""
import os, json

LOGO = "public/sop-logo.png"
BOX = "☐"  # ballot box
GOLD_HEX = "#C89B3C"; CREAM_HEX = "F5F1E6"; CREAM2_HEX = "FBF8F1"; GREY_HEX = "555555"
SCALAR = {"text", "number", "date", "time", "datetime", "select", "textarea", "checkbox"}


def load_schema(path):
    return json.load(open(path, encoding="utf-8"))


# -------- turn a schema (+ overrides) into an ordered list of layout blocks --------
def blocks_from_schema(schema, completion_as_log=None):
    """completion_as_log: optional dict {section_id: log_columns} to replace a single-entry
    section with a multi-row log table (for periodic forms)."""
    blocks = []
    for sec in schema["sections"]:
        blocks.append({"k": "section", "title": sec.get("title", ""), "desc": sec.get("description", "")})
        if completion_as_log and sec["id"] in completion_as_log:
            blocks.append({"k": "logtable", **completion_as_log[sec["id"]]})
            continue
        fields = sec.get("fields", [])
        i = 0
        while i < len(fields):
            f = fields[i]; t = f["type"]
            if t == "info":
                blocks.append({"k": "info", "text": f.get("text", "")}); i += 1
            elif t == "reference_table":
                blocks.append({"k": "reftable", "columns": f["columns"], "rows": f["rows"]}); i += 1
            elif t == "grid":
                blocks.append({"k": "grid", "field": f}); i += 1
            elif t == "signature":
                blocks.append({"k": "sig", "field": f}); i += 1
            elif t == "pass_fail":
                blocks.append({"k": "passfail", "field": f}); i += 1
            elif t in SCALAR:
                if f.get("width") == "full" or t == "textarea":
                    blocks.append({"k": "scalars", "fields": [f]}); i += 1
                else:
                    g = fields[i + 1] if i + 1 < len(fields) else None
                    if g and g["type"] in SCALAR and g.get("width") != "full" and g["type"] != "textarea":
                        blocks.append({"k": "scalars", "fields": [f, g]}); i += 2
                    else:
                        blocks.append({"k": "scalars", "fields": [f]}); i += 1
            else:
                i += 1
    return blocks


def grid_layout(field):
    rows_def = field["rows"]; cols = field["columns"]
    fixed = rows_def["mode"] == "fixed"
    lead = rows_def.get("labelHeader", "") if fixed else None
    weights = ([2.4] if fixed else []) + [c.get("width", 1) for c in cols]
    header = ([lead] if fixed else []) + [c["label"] for c in cols]
    body = []
    if fixed:
        labels = rows_def.get("labels", []); dv = rows_def.get("defaultValues", []) or []
        for idx, lab in enumerate(labels):
            d = dv[idx] if idx < len(dv) else {}
            row = [lab] + [str(d.get(c["id"], "")) for c in cols]
            body.append(row)
    else:
        for _ in range(10):
            body.append([""] * len(cols))
    return header, body, weights, fixed


# ================================ PDF ================================
def build_pdf(out, meta, blocks, landscape_page=False):
    """landscape_page: for wide register grids (FRM-004 has ten columns; portrait squeezes
    each to about half an inch, which is unreadable and unwritable)."""
    from reportlab.lib.pagesizes import letter, landscape
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.lib.utils import ImageReader

    pdfmetrics.registerFont(TTFont("Sym", "C:/Windows/Fonts/seguisym.ttf"))
    box = '<font name="Sym">☐</font>'
    GOLD = colors.HexColor(GOLD_HEX); CREAM = colors.HexColor("#" + CREAM_HEX); CREAM2 = colors.HexColor("#" + CREAM2_HEX)
    st = getSampleStyleSheet()
    base = ParagraphStyle("base", parent=st["Normal"], fontName="Helvetica", fontSize=9, leading=12)
    lbl = ParagraphStyle("lbl", parent=base, fontName="Helvetica-Bold")
    sec = ParagraphStyle("sec", parent=base, fontName="Helvetica-Bold", fontSize=11, textColor=GOLD, spaceBefore=8, spaceAfter=2)
    info = ParagraphStyle("info", parent=base, fontName="Helvetica-Oblique", fontSize=8, textColor=colors.HexColor("#" + GREY_HEX))
    cellb = ParagraphStyle("cellb", parent=base, fontName="Helvetica-Bold", fontSize=8.5)
    cell = ParagraphStyle("cell", parent=base, fontSize=8.5)
    ctr = ParagraphStyle("ctr", parent=cell, alignment=TA_CENTER)

    def P(t, s=cell): return Paragraph(str(t).replace("&", "&amp;"), s)

    page = landscape(letter) if landscape_page else letter
    doc = SimpleDocTemplate(out, pagesize=page, leftMargin=0.6 * inch, rightMargin=0.6 * inch,
                            topMargin=0.5 * inch, bottomMargin=0.6 * inch, title=f"{meta['form_no']} {meta['title']}")
    # Everything downstream sizes off W as a fraction, so the page swap is the only change needed.
    W = page[0] - 1.2 * inch; E = []
    iw, ih = ImageReader(LOGO).getSize()
    E.append(Image(LOGO, width=1.9 * inch, height=1.9 * inch * ih / iw)); E.append(Spacer(1, 6))

    # metadata header
    m = Table([
        [P("Adventure Bakery, LLC", lbl), "", P("Revision", cellb), P(meta["revision"], ctr)],
        [P("Form Title", lbl), P(meta["title"], lbl), P("Effective Date", cellb), P(meta["eff"], ctr)],
        [P("Form No.", lbl), P(meta["form_no"], base), P("Approved By", cellb), P(meta["appr"], ctr)],
        [P("SQF Reference", lbl), P(meta["sqf"], base), P("Filled By", cellb), P("", ctr)],
    ], colWidths=[W * 0.18, W * 0.42, W * 0.20, W * 0.20])
    m.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 1, colors.black), ("SPAN", (0, 0), (1, 0)),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("ALIGN", (2, 0), (3, -1), "CENTER"),
        ("ALIGN", (1, 0), (1, -1), "LEFT"), ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4)]))
    E.append(m); E.append(Spacer(1, 10))

    def scalar_value(f):
        if f["type"] == "select":
            opts = f.get("options", [])
            if opts and len(opts) <= 4 and not f.get("allowOther"):
                return P("   ".join(f"{box} {o}" for o in opts), ParagraphStyle("opt", parent=cell))
        return P("", cell)

    for b in blocks:
        k = b["k"]
        if k == "section":
            if b["title"]:
                E.append(Paragraph(b["title"].replace("&", "&amp;"), sec))
            if b["desc"]:
                E.append(Paragraph(b["desc"].replace("&", "&amp;"), info)); E.append(Spacer(1, 2))
        elif k == "info":
            E.append(Paragraph(b["text"].replace("&", "&amp;").replace("\n", "<br/>"), info)); E.append(Spacer(1, 3))
        elif k == "scalars":
            fs = b["fields"]
            if len(fs) == 1:
                f = fs[0]
                tall = f["type"] == "textarea"
                t = Table([[P(f["label"], cellb), scalar_value(f)]], colWidths=[W * 0.22, W * 0.78])
                t.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.5, colors.black), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("BACKGROUND", (0, 0), (0, 0), CREAM), ("TOPPADDING", (0, 0), (-1, -1), 14 if tall else 8),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 14 if tall else 8)]))
            else:
                f, g = fs
                t = Table([[P(f["label"], cellb), scalar_value(f), P(g["label"], cellb), scalar_value(g)]],
                          colWidths=[W * 0.16, W * 0.34, W * 0.16, W * 0.34])
                t.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.5, colors.black), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("BACKGROUND", (0, 0), (0, 0), CREAM), ("BACKGROUND", (2, 0), (2, 0), CREAM),
                    ("TOPPADDING", (0, 0), (-1, -1), 8), ("BOTTOMPADDING", (0, 0), (-1, -1), 8)]))
            E.append(t); E.append(Spacer(1, 4))
        elif k == "passfail":
            f = b["field"]; na = f.get("naAllowed")
            pf = f'Pass {box}&nbsp;&nbsp;&nbsp;Fail {box}' + (f'&nbsp;&nbsp;&nbsp;N/A {box}' if na else "")
            t = Table([[P(f["label"], cellb), Paragraph(pf, ParagraphStyle("pf", parent=cell, fontSize=10))]],
                      colWidths=[W * 0.62, W * 0.38])
            t.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.5, colors.black), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("BACKGROUND", (0, 0), (0, 0), CREAM), ("TOPPADDING", (0, 0), (-1, -1), 8), ("BOTTOMPADDING", (0, 0), (-1, -1), 8)]))
            E.append(t); E.append(Spacer(1, 4))
        elif k == "sig":
            f = b["field"]
            t = Table([[P(f["label"], cellb), P("", cell), P("Date", cellb), P("", cell)]],
                      colWidths=[W * 0.26, W * 0.40, W * 0.10, W * 0.24])
            t.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.5, colors.black), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("BACKGROUND", (0, 0), (0, 0), CREAM), ("BACKGROUND", (2, 0), (2, 0), CREAM),
                ("TOPPADDING", (0, 0), (-1, -1), 12), ("BOTTOMPADDING", (0, 0), (-1, -1), 12)]))
            E.append(t)
            if f.get("statement"):
                E.append(Paragraph('"' + f["statement"].replace("&", "&amp;") + '"', info))
            E.append(Spacer(1, 4))
        elif k == "reftable":
            data = [[P(c, cellb) for c in b["columns"]]] + [[P(c, cell) for c in r] for r in b["rows"]]
            n = len(b["columns"]); t = Table(data, colWidths=[W / n] * n, repeatRows=1)
            gs = [("GRID", (0, 0), (-1, -1), 0.5, colors.black), ("BACKGROUND", (0, 0), (-1, 0), CREAM),
                  ("VALIGN", (0, 0), (-1, -1), "TOP"), ("TOPPADDING", (0, 1), (-1, -1), 4), ("BOTTOMPADDING", (0, 1), (-1, -1), 4)]
            t.setStyle(TableStyle(gs)); E.append(t); E.append(Spacer(1, 6))
        elif k == "grid":
            header, body, weights, fixed = grid_layout(b["field"])
            tot = sum(weights); cw = [W * w / tot for w in weights]
            # Wide grids: reportlab's default 6pt side padding costs 12pt of every column, and at
            # 8.5pt a one-weight column cannot hold "Sanitation" - reportlab then splits the word
            # itself ("Sanitatio / n SOP"). Tighten the padding and the type instead of distorting
            # the schema's column weights, which have to serve the screen too.
            wide = len(header) >= 8
            hstyle = ParagraphStyle("cellbw", parent=cellb, fontSize=7.5, leading=9) if wide else cellb
            cstyle = ParagraphStyle("cellw", parent=cell, fontSize=7.5, leading=9) if wide else cell
            data = [[P(h, hstyle) for h in header]] + [[P(c, cstyle) for c in r] for r in body]
            t = Table(data, colWidths=cw, repeatRows=1)
            gs = [("GRID", (0, 0), (-1, -1), 0.5, colors.black), ("BACKGROUND", (0, 0), (-1, 0), CREAM),
                  ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("TOPPADDING", (0, 1), (-1, -1), 9), ("BOTTOMPADDING", (0, 1), (-1, -1), 9)]
            if wide:
                # Row padding also comes down: a wide register's rows arrive pre-filled, so only
                # the couple of genuinely blank cells need writing room, and 9pt of padding per row
                # was pushing the closing note onto a page of its own.
                gs += [("LEFTPADDING", (0, 0), (-1, -1), 2), ("RIGHTPADDING", (0, 0), (-1, -1), 2),
                       ("TOPPADDING", (0, 1), (-1, -1), 6), ("BOTTOMPADDING", (0, 1), (-1, -1), 6)]
            if fixed:
                for i in range(1, len(data)):
                    gs.append(("BACKGROUND", (0, i), (0, i), CREAM2))
            t.setStyle(TableStyle(gs)); E.append(t); E.append(Spacer(1, 6))
        elif k == "logtable":
            cols = b["columns"]; nrows = b.get("nrows", 12); weights = b.get("weights", [1] * len(cols))
            tot = sum(weights); cw = [W * w / tot for w in weights]
            data = [[P(c, cellb) for c in cols]] + [[P("", cell) for _ in cols] for _ in range(nrows)]
            t = Table(data, colWidths=cw, repeatRows=1)
            t.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.5, colors.black), ("BACKGROUND", (0, 0), (-1, 0), CREAM),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"), ("TOPPADDING", (0, 1), (-1, -1), 10), ("BOTTOMPADDING", (0, 1), (-1, -1), 10)]))
            E.append(t); E.append(Spacer(1, 6))

    E.append(Spacer(1, 8)); E.append(Paragraph(meta["footer"], info))
    doc.build(E); print("wrote", out, os.path.getsize(out), "bytes")


# ================================ DOCX ================================
def build_docx(out, meta, blocks, landscape_page=False):
    """landscape_page: see build_pdf. python-docx does not swap the page dimensions when the
    orientation changes, so they are swapped by hand."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH as AL
    from docx.enum.table import WD_ALIGN_VERTICAL as VA
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    GOLD = RGBColor(0xC8, 0x9B, 0x3C); GREY = RGBColor(0x55, 0x55, 0x55)
    d = Document(); s0 = d.sections[0]
    if landscape_page:
        s0.orientation = WD_ORIENT.LANDSCAPE
        s0.page_width, s0.page_height = s0.page_height, s0.page_width
    s0.top_margin = Inches(0.5); s0.bottom_margin = Inches(0.6); s0.left_margin = Inches(0.6); s0.right_margin = Inches(0.6)
    d.styles["Normal"].font.name = "Calibri"; d.styles["Normal"].font.size = Pt(9)
    PAGEW = s0.page_width.inches - 1.2  # inches usable (7.3 portrait, 9.8 landscape)

    def shade(c, hexc):
        tcPr = c._tc.get_or_add_tcPr(); sh = OxmlElement("w:shd")
        sh.set(qn("w:val"), "clear"); sh.set(qn("w:fill"), hexc); tcPr.append(sh)

    def para(c, text, bold=False, size=9, color=None, align=None, italic=False, font=None):
        p = c.paragraphs[0] if (c.paragraphs and not c.paragraphs[0].runs) else c.add_paragraph()
        if align is not None: p.alignment = align
        for j, seg in enumerate(str(text).split("\n")):
            if j: p = c.add_paragraph()
            r = p.add_run(seg); r.bold = bold; r.italic = italic; r.font.size = Pt(size)
            if color is not None: r.font.color.rgb = color
            if font: r.font.name = font
        return p

    def widths(tbl, ws):
        # Call sites pass absolute inches that add up to the portrait usable width; normalising
        # by the total makes them relative, so landscape scales them and portrait is unchanged.
        scale = PAGEW / sum(ws)
        for row in tbl.rows:
            for i, c in enumerate(row.cells):
                c.width = Inches(ws[i] * scale); c.vertical_alignment = VA.CENTER

    def checkbox_run(p, size=11):
        r = p.add_run(BOX); r.font.name = "Segoe UI Symbol"; r.font.size = Pt(size); return r

    p = d.add_paragraph(); p.alignment = AL.CENTER; p.add_run().add_picture(LOGO, width=Inches(1.9))

    m = d.add_table(rows=4, cols=4); m.style = "Table Grid"; m.alignment = 1
    widths(m, [1.3, 3.0, 1.5, 1.5])
    a = m.cell(0, 0); a.merge(m.cell(0, 1)); para(a, "Adventure Bakery, LLC", bold=True)
    para(m.cell(0, 2), "Revision", bold=True, align=AL.CENTER); para(m.cell(0, 3), meta["revision"], align=AL.CENTER)
    para(m.cell(1, 0), "Form Title", bold=True); para(m.cell(1, 1), meta["title"], bold=True)
    para(m.cell(1, 2), "Effective Date", bold=True, align=AL.CENTER); para(m.cell(1, 3), meta["eff"], align=AL.CENTER)
    para(m.cell(2, 0), "Form No.", bold=True); para(m.cell(2, 1), meta["form_no"])
    para(m.cell(2, 2), "Approved By", bold=True, align=AL.CENTER); para(m.cell(2, 3), meta["appr"], align=AL.CENTER)
    para(m.cell(3, 0), "SQF Reference", bold=True); para(m.cell(3, 1), meta["sqf"])
    para(m.cell(3, 2), "Filled By", bold=True, align=AL.CENTER); para(m.cell(3, 3), "")
    d.add_paragraph()

    def sec_head(t):
        h = d.add_paragraph(); r = h.add_run(t); r.bold = True; r.font.size = Pt(12); r.font.color.rgb = GOLD

    def opt_cell(c, f):
        if f["type"] == "select":
            opts = f.get("options", [])
            if opts and len(opts) <= 4 and not f.get("allowOther"):
                p = c.paragraphs[0]
                for k, o in enumerate(opts):
                    if k: p.add_run("   ")
                    checkbox_run(p, 11); p.add_run(" " + o)
                return
        c.add_paragraph()

    for b in blocks:
        k = b["k"]
        if k == "section":
            if b["title"]: sec_head(b["title"])
            if b["desc"]:
                q = d.add_paragraph(); r = q.add_run(b["desc"]); r.italic = True; r.font.size = Pt(8); r.font.color.rgb = GREY
        elif k == "info":
            q = d.add_paragraph(); r = q.add_run(b["text"]); r.italic = True; r.font.size = Pt(8); r.font.color.rgb = GREY
        elif k == "scalars":
            fs = b["fields"]
            if len(fs) == 1:
                f = fs[0]; t = d.add_table(rows=1, cols=2); t.style = "Table Grid"; widths(t, [1.7, 5.6])
                para(t.cell(0, 0), f["label"], bold=True); shade(t.cell(0, 0), CREAM_HEX); opt_cell(t.cell(0, 1), f)
                if f["type"] == "textarea": t.cell(0, 1).add_paragraph()
            else:
                f, g = fs; t = d.add_table(rows=1, cols=4); t.style = "Table Grid"; widths(t, [1.3, 2.35, 1.3, 2.35])
                para(t.cell(0, 0), f["label"], bold=True); shade(t.cell(0, 0), CREAM_HEX); opt_cell(t.cell(0, 1), f)
                para(t.cell(0, 2), g["label"], bold=True); shade(t.cell(0, 2), CREAM_HEX); opt_cell(t.cell(0, 3), g)
            d.add_paragraph()
        elif k == "passfail":
            f = b["field"]; t = d.add_table(rows=1, cols=2); t.style = "Table Grid"; widths(t, [4.5, 2.8])
            para(t.cell(0, 0), f["label"], bold=True); shade(t.cell(0, 0), CREAM_HEX)
            p = t.cell(0, 1).paragraphs[0]; p.add_run("Pass "); checkbox_run(p); p.add_run("    Fail "); checkbox_run(p)
            if f.get("naAllowed"): p.add_run("    N/A "); checkbox_run(p)
            d.add_paragraph()
        elif k == "sig":
            f = b["field"]; t = d.add_table(rows=1, cols=4); t.style = "Table Grid"; widths(t, [1.9, 2.9, 0.7, 1.8])
            para(t.cell(0, 0), f["label"], bold=True); shade(t.cell(0, 0), CREAM_HEX); t.cell(0, 1).add_paragraph()
            para(t.cell(0, 2), "Date", bold=True); shade(t.cell(0, 2), CREAM_HEX); t.cell(0, 3).add_paragraph()
            if f.get("statement"):
                q = d.add_paragraph(); r = q.add_run('"' + f["statement"] + '"'); r.italic = True; r.font.size = Pt(8); r.font.color.rgb = GREY
            d.add_paragraph()
        elif k == "reftable":
            cols = b["columns"]; rows = b["rows"]; n = len(cols)
            t = d.add_table(rows=1 + len(rows), cols=n); t.style = "Table Grid"; widths(t, [PAGEW / n] * n)
            for i, c in enumerate(cols):
                para(t.rows[0].cells[i], c, bold=True); shade(t.rows[0].cells[i], CREAM_HEX)
            for ri, row in enumerate(rows, 1):
                for i, val in enumerate(row):
                    t.rows[ri].cells[i].vertical_alignment = VA.TOP; para(t.rows[ri].cells[i], val, size=8)
            d.add_paragraph()
        elif k == "grid":
            header, body, weights, fixed = grid_layout(b["field"]); tot = sum(weights)
            ws = [PAGEW * w / tot for w in weights]
            gfs = 7.5 if len(header) >= 8 else 8.5  # match the PDF's wide-grid type size
            t = d.add_table(rows=1 + len(body), cols=len(header)); t.style = "Table Grid"; widths(t, ws)
            for i, h in enumerate(header):
                para(t.rows[0].cells[i], h, bold=True, size=gfs); shade(t.rows[0].cells[i], CREAM_HEX)
            for ri, row in enumerate(body, 1):
                for i, val in enumerate(row):
                    cc = t.rows[ri].cells[i]; para(cc, val, size=gfs)
                    if fixed and i == 0: shade(cc, CREAM2_HEX)
                    cc.add_paragraph()
            d.add_paragraph()
        elif k == "logtable":
            cols = b["columns"]; nrows = b.get("nrows", 12); weights = b.get("weights", [1] * len(cols)); tot = sum(weights)
            ws = [PAGEW * w / tot for w in weights]
            t = d.add_table(rows=1 + nrows, cols=len(cols)); t.style = "Table Grid"; widths(t, ws)
            for i, c in enumerate(cols):
                para(t.rows[0].cells[i], c, bold=True, size=8.5); shade(t.rows[0].cells[i], CREAM_HEX)
            for ri in range(1, nrows + 1):
                for i in range(len(cols)):
                    t.rows[ri].cells[i].add_paragraph()
            d.add_paragraph()

    foot = d.add_paragraph(); r = foot.add_run(meta["footer"]); r.italic = True; r.font.size = Pt(8); r.font.color.rgb = GREY
    d.save(out); print("wrote", out, os.path.getsize(out), "bytes")


# ================================ drive ================================
FOOT = "Adventure Bakery, LLC · Confidential · {no} · Handwritten entries are photographed and entered into the digital form"

if __name__ == "__main__":
    # FRM-903 — one sheet per production day
    s903 = load_schema("sop-drafts/FRM-903-consolidated-daily-record-schema.json")
    # Document-control values mirror the live sop_documents row for FRM-903, read
    # 2026-08-21: rev v2, effective 2026-08-21, active. sqf_reference is the completed
    # five-clause list applied by migration 20260821000001 (one clause per section).
    # Re-check all of these against the row whenever the schema changes - a blank printed
    # with a revision the quality system does not have is a document-control finding, not
    # a cosmetic one.
    meta903 = {"form_no": "FRM-903", "title": "Daily Sanitation, Pre-Operation & Release Record",
               "revision": "v2", "eff": "2026-08-21", "appr": "GJM",
               "sqf": "11.2.5.1, 11.2.5.3, 11.2.5.7, 11.2.5.9, 11.7.3.4",
               "footer": FOOT.format(no="FRM-903")}
    b903 = blocks_from_schema(s903)
    build_pdf("sop-drafts/FRM-903-blank.pdf", meta903, b903)
    build_docx("sop-drafts/FRM-903-blank.docx", meta903, b903)

    # FRM-901 — schedule (printed) + periodic completion LOG (each row -> one digital entry)
    s901 = load_schema("sop-drafts/FRM-901-reslim-schema.json")
    meta901 = {"form_no": "FRM-901", "title": "Master Sanitation Schedule",
               "revision": "A", "eff": "2026-06-01", "appr": "GJM",
               "sqf": "11.2.5.1, 11.2.5.8, 11.2.5.9", "footer": FOOT.format(no="FRM-901")}
    completion_log = {"completion": {
        "columns": ["Task / area cleaned", "Frequency", "Date completed", "Method / chemical / comments", "Completed by", "Verified by / date"],
        "weights": [2.4, 1.1, 1.1, 2.3, 1.2, 1.4], "nrows": 12}}
    b901 = blocks_from_schema(s901, completion_as_log=completion_log)
    build_pdf("sop-drafts/FRM-901-blank.pdf", meta901, b901)
    build_docx("sop-drafts/FRM-901-blank.docx", meta901, b901)

    # FRM-004 — equipment register. LANDSCAPE: ten columns (leading label + nine) do not fit
    # portrait; each would land at about half an inch, too narrow to read or write in.
    #
    # The live row is status='draft' with no effective_date and no approved_by (read 2026-08-22),
    # so the blank says so on its face rather than printing a blank date box that reads as an
    # oversight. Re-stamp these here from the row when the activation migration runs - a blank
    # circulating as if approved when it is not is the document-control failure this avoids.
    s004 = load_schema("sop-drafts/FRM-004-equipment-register-schema.json")
    meta004 = {"form_no": "FRM-004", "title": "Equipment Register",
               "revision": "New", "eff": "Draft — not yet effective", "appr": "—",
               "sqf": "11.2.1.2",
               "footer": FOOT.format(no="FRM-004") + " · DRAFT — not approved for use"}
    b004 = blocks_from_schema(s004)
    build_pdf("sop-drafts/FRM-004-blank.pdf", meta004, b004, landscape_page=True)
    build_docx("sop-drafts/FRM-004-blank.docx", meta004, b004, landscape_page=True)
